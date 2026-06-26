"""bid_engine.py -- AI-powered government tender bid document generator.

Flow:
  1. extract_tender(content, mime_type)          -- parse tender PDF/image -> structured dict
  2. readiness_check(tender, profile, texts)     -- what's met, missing, needs attention
  3. generate_bid_content(tender, profile, texts)-- Claude writes the full bid JSON
  4. build_docx(bid_json, tender, profile)       -- compile into a Word document (bytes)
"""
from __future__ import annotations

import io
import json
import logging
import os
from datetime import date

import core

logger = logging.getLogger(__name__)

# Pinned to gemini-2.5-flash: the "-latest" alias returns 503 and 2.0-flash 429s
# for AQ.-format keys; 2.5-flash authenticates and serves reliably (verified 200).
GEMINI_MODEL = "gemini-2.5-flash"


# ---------------------------------------------------------------------------
# AI helper — tries Gemini first, falls back to Claude
# ---------------------------------------------------------------------------

def _call_ai(prompt: str, doc_bytes: bytes | None = None,
             mime_type: str = "application/pdf", max_tokens: int = 3000) -> dict | str | None:
    """Call Gemini (preferred) or Claude to get a JSON response."""
    import base64

    gemini_key    = os.getenv("GEMINI_API_KEY")
    anthropic_key = os.getenv("ANTHROPIC_API_KEY")

    if not gemini_key and not anthropic_key:
        core.record_ai_error("no api key configured")
        return None

    full_prompt = prompt + "\n\nReturn ONLY valid raw JSON. No markdown fences."

    # ---- Gemini path (direct REST + X-goog-api-key header) ----
    # Direct REST avoids the google-genai SDK's mishandling of the newer "AQ."
    # API-key format (which the SDK sends as an OAuth token -> 401).
    if gemini_key:
        text = ""
        try:
            import requests
            parts: list = []
            if doc_bytes:
                actual_mime = mime_type
                if "pdf" not in mime_type.lower() and "image" not in mime_type.lower():
                    actual_mime = "application/pdf"
                parts.append({"inline_data": {
                    "mime_type": actual_mime,
                    "data": base64.b64encode(doc_bytes).decode(),
                }})
            parts.append({"text": full_prompt})

            import time as _time
            resp = None
            for _attempt in range(3):
                resp = requests.post(
                    f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent",
                    headers={"Content-Type": "application/json", "X-goog-api-key": gemini_key},
                    # thinkingBudget:0 -> far fewer tokens/quota per call, no quality
                    # loss on these heavily-structured prompts.
                    json={"contents": [{"parts": parts}],
                          "generationConfig": {"responseMimeType": "application/json",
                                               "thinkingConfig": {"thinkingBudget": 0}}},
                    timeout=120,
                )
                # Gemini's free tier intermittently 503/429s ("high demand") — retry
                # with backoff so a transient spike doesn't fail the whole draft.
                if resp.status_code not in (503, 429):
                    break
                if _attempt < 2:
                    _time.sleep(1.0 * (_attempt + 1))
            resp.raise_for_status()
            _parts = resp.json()["candidates"][0]["content"]["parts"]
            text = "".join(p.get("text", "") for p in _parts).strip()
            text = text.removeprefix("```json").removeprefix("```").removesuffix("```").strip()
            result = json.loads(text)
            core.clear_ai_error()
            return result
        except json.JSONDecodeError:
            core.clear_ai_error()
            return text  # model returned prose, not JSON — hand it back as-is
        except Exception as exc:
            logger.warning("Gemini API error: %s", exc)
            core.record_ai_error(exc)
            if not anthropic_key:
                return None
            # fall through to Claude

    # ---- Claude fallback ----
    if anthropic_key:
        try:
            import anthropic
            client  = anthropic.Anthropic(api_key=anthropic_key)
            content: list = []

            if doc_bytes:
                enc = base64.b64encode(doc_bytes).decode()
                if "pdf" in mime_type.lower():
                    content.append({
                        "type": "document",
                        "source": {"type": "base64", "media_type": "application/pdf", "data": enc},
                    })
                else:
                    actual = ("image/jpeg" if any(x in mime_type for x in ("jpg","jpeg")) else "image/png")
                    content.append({
                        "type": "image",
                        "source": {"type": "base64", "media_type": actual, "data": enc},
                    })

            content.append({"type": "text", "text": full_prompt})
            msg  = client.messages.create(
                model="claude-sonnet-4-6", max_tokens=max_tokens,
                messages=[{"role": "user", "content": content}],
            )
            text = "".join(b.text for b in msg.content if getattr(b,"type","") == "text").strip()
            text = text.removeprefix("```json").removeprefix("```").removesuffix("```").strip()
            result = json.loads(text)
            core.clear_ai_error()
            return result
        except json.JSONDecodeError:
            core.clear_ai_error()
            return text  # type: ignore[possibly-undefined]
        except Exception as exc:
            logger.warning("Claude API error: %s", exc)
            core.record_ai_error(exc)
            return None

    return None


# keep old name as alias so nothing else breaks
def _claude(prompt: str, doc_bytes: bytes | None = None,
            mime_type: str = "application/pdf", max_tokens: int = 3000) -> dict | str | None:
    return _call_ai(prompt, doc_bytes, mime_type, max_tokens)


# ---------------------------------------------------------------------------
# Step 1 -- Extract tender details from uploaded document
# ---------------------------------------------------------------------------

def extract_tender(content: bytes, mime_type: str = "application/pdf") -> dict:
    """Parse a tender notice PDF/image and return a structured dict.

    Falls back to an empty skeleton when ANTHROPIC_API_KEY is not set.
    """
    PROMPT = """You are analyzing an Indian government tender / NIT (Notice Inviting Tender) document.
Extract every key field precisely. Use null for any field not found.

Return ONLY valid JSON:
{
  "title": "Full tender title",
  "tender_no": "Tender or NIT reference number",
  "organization": "Full name of issuing department/authority",
  "state": "State name",
  "district": "District if mentioned, else null",
  "category": "Type of work (Civil Infrastructure / Electrical / IT / Medical / Water / Transport etc.)",
  "value_text": "Estimated cost as written (e.g. Rs 45.50 Lakhs)",
  "value_lakhs": null,
  "deadline": "Submission closing date as YYYY-MM-DD if possible, else as written",
  "emd": "Earnest money deposit amount with currency",
  "contractor_class": "Required class (Class A/B/C/D) or null",
  "required_turnover_lakhs": null,
  "required_experience_years": null,
  "scope_of_work": "Clear 2-3 sentence description of work to be performed",
  "required_documents": ["GST registration", "PAN card", "Experience certificate", "..."],
  "eligibility_criteria": "Full eligibility / qualification criteria text",
  "submission_mode": "Online / Offline / Both",
  "contact_office": "Office / email / phone for tender queries"
}"""

    result = _claude(PROMPT, content, mime_type, max_tokens=1500)

    if not isinstance(result, dict):
        return {
            "title": "Tender (add ANTHROPIC_API_KEY to enable AI extraction)",
            "tender_no": None, "organization": None, "state": None,
            "district": None, "category": None, "value_text": None,
            "value_lakhs": None, "deadline": None, "emd": None,
            "contractor_class": None, "required_turnover_lakhs": None,
            "required_experience_years": None,
            "scope_of_work": "Could not extract — ANTHROPIC_API_KEY required.",
            "required_documents": [], "eligibility_criteria": None,
            "submission_mode": None, "contact_office": None,
            "_extraction_failed": True,
        }

    if result.get("value_text") and not result.get("value_lakhs"):
        result["value_lakhs"] = core.parse_value_to_lakhs(result["value_text"])

    return result


# ---------------------------------------------------------------------------
# Step 2 -- Readiness check (profile + vault docs vs tender requirements)
# ---------------------------------------------------------------------------

_DOC_KEYWORDS = {
    "GST Registration":             ["gst"],
    "PAN Card":                     ["pan "],
    "Contractor Registration":      ["registration", "license", "licence", "empanel"],
    "EMD / Bank Guarantee":         ["emd", "bank guarantee", "earnest money", "demand draft"],
    "Turnover / Financial Statement":["turnover", "balance sheet", "audited accounts"],
    "Experience / Completion Cert": ["experience", "completion certificate", "similar work"],
    "ISO Certification":            ["iso "],
    "DGMS / Safety Certification":  ["dgms", "safety pass"],
}


def readiness_check(tender: dict, profile: dict, vault_texts: list[str]) -> dict:
    """Compare tender requirements against profile + vault documents.

    Returns dict with keys: met, missing, warnings, overall_pct, ready.
    """
    combined = " ".join(vault_texts).lower()
    p_class    = core._rank(profile.get("contractor_class", ""))

    def _sf(v, d=0.0):
        try:
            f = float(v or d)
            return d if f != f else f
        except (TypeError, ValueError):
            return d

    p_turnover = _sf(profile.get("turnover_lakhs"))
    p_exp      = int(_sf(profile.get("experience_years")))

    met, missing, warnings = [], [], []

    # Contractor class
    req_cls = tender.get("contractor_class")
    if req_cls and core._rank(req_cls) > 0:
        if p_class >= core._rank(req_cls):
            met.append(f"Contractor class: you hold {profile.get('contractor_class')} (need {req_cls})")
        else:
            missing.append(f"Class {req_cls} required — you have {profile.get('contractor_class') or 'not set'}")

    # Turnover
    req_to = tender.get("required_turnover_lakhs")
    if not req_to and tender.get("value_lakhs"):
        try:
            req_to = round(float(tender["value_lakhs"]) * 0.3, 1)
        except (TypeError, ValueError):
            pass
    if req_to:
        req_to = float(req_to)
        if p_turnover >= req_to:
            met.append(f"Turnover: Rs {p_turnover:.0f}L (need ~Rs {req_to:.0f}L)")
        elif p_turnover > 0:
            missing.append(f"Turnover may be short: you have Rs {p_turnover:.0f}L, need ~Rs {req_to:.0f}L")
        else:
            warnings.append("Turnover not set in Profile — update it to verify eligibility")

    # Experience
    req_exp = tender.get("required_experience_years")
    if not req_exp:
        req_exp = core._exp_years(tender.get("eligibility_criteria","") or "")
    if req_exp:
        req_exp = int(req_exp)
        if p_exp >= req_exp:
            met.append(f"Experience: {p_exp} years (need {req_exp}+)")
        else:
            missing.append(f"Experience: {req_exp}+ years needed, profile shows {p_exp} years")

    # Required documents listed in tender
    for doc in (tender.get("required_documents") or []):
        doc_l  = doc.lower()
        in_vault = any(word in combined for word in doc_l.split()[:3]) if combined else False
        if in_vault:
            met.append(f"Document found in vault: {doc}")
        else:
            missing.append(f"Required document not in vault: {doc}")

    # Scan vault text for common doc types
    if combined:
        for label, kws in _DOC_KEYWORDS.items():
            if any(kw in combined for kw in kws):
                if not any(label.lower() in m.lower() for m in met):
                    met.append(f"Found in uploaded documents: {label}")

    # Deadline warning
    deadline = core.parse_date(tender.get("deadline"))
    if deadline:
        dl = (deadline - date.today()).days
        if dl < 0:
            warnings.append(f"Deadline has already passed ({tender.get('deadline')})")
        elif dl <= 3:
            warnings.append(f"Only {dl} day(s) left — expedite immediately")
        elif dl <= 7:
            warnings.append(f"Deadline in {dl} days — prioritize preparation")

    # Value vs turnover risk
    val = tender.get("value_lakhs")
    if val and p_turnover > 0:
        try:
            if float(val) > p_turnover * 2:
                warnings.append(
                    f"Tender value Rs {float(val):.0f}L exceeds 2x your turnover — may fail eligibility"
                )
        except (TypeError, ValueError):
            pass

    total = len(met) + len(missing)
    pct   = round(100 * len(met) / total) if total else 50

    return {"met": met, "missing": missing, "warnings": warnings,
            "overall_pct": pct, "ready": len(missing) == 0}


# ---------------------------------------------------------------------------
# Step 3 -- Generate full bid content with Claude
# ---------------------------------------------------------------------------

def generate_bid_content(tender: dict, profile: dict,
                          vault_texts: list[str],
                          language: str = "en") -> dict | None:
    """Call Claude to generate the complete bid document content as JSON.

    language: 'en' (English) or 'hi' (Hindi / Devanagari). The JSON KEYS stay
    English (build_docx relies on them); only the human-readable VALUES are
    written in the chosen language.
    """
    cname    = profile.get("company_name") or "Our Firm"
    cls      = profile.get("contractor_class") or "Class C"
    exp      = profile.get("experience_years") or 0
    turnover = profile.get("turnover_lakhs") or 0
    sectors  = ", ".join(profile.get("sectors") or ["General Construction"])
    states   = ", ".join(profile.get("states") or ["Chhattisgarh"])
    vault_summary = (" ".join(vault_texts)[:1500]) if vault_texts else "No additional documents uploaded."

    lang_directive = ""
    if str(language).lower().startswith("hi"):
        lang_directive = (
            "\nIMPORTANT LANGUAGE RULE: Write every human-readable text VALUE in "
            "formal Hindi (Devanagari script) suitable for an Indian government "
            "tender. Keep all JSON keys, the structure, proper nouns, numbers, "
            "currency and the word 'EMD/GST/PAN' in English/as-is. Do NOT translate "
            "the JSON field names.\n")

    PROMPT = f"""You are an expert Indian government tender bid writing consultant.
Create a complete, professional, and highly personalized bid package.{lang_directive}

TENDER DETAILS:
Title: {tender.get('title','—')}
Tender No: {tender.get('tender_no','—')}
Organization: {tender.get('organization','—')}
Location: {tender.get('state','—')} / {tender.get('district','—')}
Category: {tender.get('category','—')}
Value: {tender.get('value_text','—')}
Deadline: {tender.get('deadline','—')}
EMD: {tender.get('emd','—')}
Scope: {tender.get('scope_of_work','—')}
Eligibility: {tender.get('eligibility_criteria','—')}
Required docs: {', '.join(tender.get('required_documents') or [])}

CONTRACTOR PROFILE:
Company: {cname}
Class: {cls}   Experience: {exp} years   Turnover: Rs {turnover} Lakhs
Sectors: {sectors}   States: {states}
Documents in vault: {vault_summary}

Write a COMPLETE bid package. Every section must be specific to this tender — no generic filler.
Return ONLY valid JSON (no markdown fences):
{{
  "cover_letter": {{
    "to": "The Tender Committee / Authority, {tender.get('organization','—')}",
    "subject": "Submission of Technical and Financial Bid for: {(tender.get('title') or '')[:80]}",
    "ref": "Ref: NIT/Tender No. {tender.get('tender_no','—')}",
    "salutation": "Respected Sir/Madam,",
    "body_paragraphs": [
      "Opening paragraph with formal expression of interest",
      "Company credentials paragraph mentioning {exp} years experience and Rs {turnover}L turnover",
      "Specific relevance to this tender's work scope",
      "Commitment to quality, timelines and compliance paragraph"
    ],
    "closing": "Closing paragraph requesting favorable consideration"
  }},
  "company_profile": {{
    "overview": "3-4 sentence company overview highlighting {cname}'s specialization in {sectors} with {exp} years experience",
    "key_strengths": [
      "Strength 1 specific to this tender type",
      "Strength 2",
      "Strength 3",
      "Strength 4"
    ]
  }},
  "technical_proposal": {{
    "scope_understanding": "Paragraph showing deep understanding of the specific work described in this tender",
    "methodology": [
      "Phase 1: Planning and mobilization (description)",
      "Phase 2: Execution step (description)",
      "Phase 3: Quality verification (description)",
      "Phase 4: Testing, commissioning and handover (description)"
    ],
    "team_structure": "Description of proposed team: Project Manager, Site Engineer, Quality Inspector, etc.",
    "quality_assurance": "Paragraph on QA/QC measures specific to this type of work",
    "timeline_overview": "Realistic timeline narrative for executing this specific scope of work"
  }},
  "compliance_matrix": [
    {{"requirement": "Contractor Class", "our_response": "{cls} registration held", "status": "Compliant"}},
    {{"requirement": "Annual Turnover", "our_response": "Rs {turnover} Lakhs (CA certified)", "status": "Compliant"}},
    {{"requirement": "Years of Experience", "our_response": "{exp} years in {sectors}", "status": "Compliant"}},
    {{"requirement": "EMD", "our_response": "DD/Bank Guarantee enclosed as per requirement", "status": "Compliant"}},
    {{"requirement": "GST Registration", "our_response": "Valid GST certificate enclosed", "status": "Compliant"}}
  ],
  "document_checklist": [
    {{"document": "Duly filled tender form", "status": "Enclosed"}},
    {{"document": "EMD / Earnest Money", "status": "DD enclosed"}},
    {{"document": "GST Registration Certificate", "status": "Self-attested copy enclosed"}},
    {{"document": "PAN Card copy", "status": "Self-attested copy enclosed"}},
    {{"document": "Contractor Registration Certificate ({cls})", "status": "Self-attested copy enclosed"}},
    {{"document": "Annual Turnover Certificate (CA certified)", "status": "Enclosed for last 3 years"}},
    {{"document": "Work Experience / Completion Certificates", "status": "Copies enclosed"}},
    {{"document": "Affidavit of non-blacklisting", "status": "Notarized copy enclosed"}}
  ],
  "declaration": "I/We, the undersigned, hereby declare that the information furnished above is true and correct to the best of my/our knowledge and belief. I/We agree to abide by all the terms and conditions of the tender document and undertake to execute the work as per specifications if awarded the contract."
}}"""

    result = _claude(PROMPT, max_tokens=4000)
    return _coerce_to_dict(result)


def _coerce_to_dict(result) -> dict | None:
    """Rescue a bid payload into a dict.

    _claude returns a dict on success, but hands back a raw string when the model
    wraps its JSON in markdown fences or appends conversational text. Strip any
    fences and attempt an explicit parse; as a last resort, extract the outermost
    {...} block. Only return None if nothing usable can be recovered."""
    if isinstance(result, dict):
        return result
    if isinstance(result, str):
        cleaned = (result.strip()
                   .removeprefix("```json").removeprefix("```")
                   .removesuffix("```").strip())
        try:
            parsed = json.loads(cleaned)
            return parsed if isinstance(parsed, dict) else None
        except (json.JSONDecodeError, ValueError):
            import re as _re
            m = _re.search(r"\{.*\}", cleaned, _re.DOTALL)
            if m:
                try:
                    parsed = json.loads(m.group(0))
                    return parsed if isinstance(parsed, dict) else None
                except (json.JSONDecodeError, ValueError):
                    pass
    return None


# ---------------------------------------------------------------------------
# Step 4 -- Compile DOCX
# ---------------------------------------------------------------------------

# Static DOCX labels (the AI-written body text is already in the chosen language;
# these are the fixed headings/labels the template prints around it).
_DOCX_LABELS = {
    "en": {
        "bid_document": "BID DOCUMENT", "tender_no": "Tender No", "issued_by": "Issued by",
        "submitted_by": "Submitted by", "date": "Date", "cover_letter": "COVER LETTER",
        "to_authority": "To The Tendering Authority,", "salutation": "Respected Sir/Madam,",
        "yours_faithfully": "Yours faithfully,", "authorized_signatory": "Authorized Signatory",
        "sec_company": "SECTION 1 — COMPANY PROFILE", "firm_details": "Firm Details",
        "company_name": "Company Name", "contractor_class": "Contractor Class",
        "annual_turnover": "Annual Turnover", "experience": "Experience",
        "sectors": "Sectors", "states": "States", "key_strengths": "Key Strengths",
        "sec_technical": "SECTION 2 — TECHNICAL PROPOSAL",
        "t_scope": "2.1  Scope Understanding", "t_method": "2.2  Execution Methodology",
        "t_team": "2.3  Team Structure", "t_qa": "2.4  Quality Assurance",
        "t_timeline": "2.5  Timeline", "sec_compliance": "SECTION 3 — COMPLIANCE MATRIX",
        "compliance_intro": "Confirmation of compliance with each tender requirement:",
        "h_requirement": "Requirement", "h_response": "Our Response", "h_status": "Status",
        "sec_checklist": "SECTION 4 — DOCUMENT CHECKLIST",
        "checklist_intro": "Documents enclosed with this bid submission:",
        "h_sr": "Sr.", "h_document": "Document", "declaration": "DECLARATION",
        "years": "years", "signatory_seal": "Authorized Signatory  |  Seal",
    },
    "hi": {
        "bid_document": "बिड दस्तावेज़", "tender_no": "टेंडर क्रमांक", "issued_by": "जारीकर्ता",
        "submitted_by": "प्रस्तुतकर्ता", "date": "दिनांक", "cover_letter": "कवर लेटर",
        "to_authority": "सेवा में, निविदा प्राधिकारी,", "salutation": "महोदय/महोदया,",
        "yours_faithfully": "भवदीय,", "authorized_signatory": "अधिकृत हस्ताक्षरकर्ता",
        "sec_company": "खंड 1 — कंपनी प्रोफ़ाइल", "firm_details": "फर्म विवरण",
        "company_name": "कंपनी का नाम", "contractor_class": "ठेकेदार श्रेणी",
        "annual_turnover": "वार्षिक टर्नओवर", "experience": "अनुभव",
        "sectors": "क्षेत्र", "states": "राज्य", "key_strengths": "मुख्य ताकतें",
        "sec_technical": "खंड 2 — तकनीकी प्रस्ताव",
        "t_scope": "2.1  कार्य की समझ", "t_method": "2.2  कार्य-निष्पादन पद्धति",
        "t_team": "2.3  टीम संरचना", "t_qa": "2.4  गुणवत्ता आश्वासन",
        "t_timeline": "2.5  समय-सीमा", "sec_compliance": "खंड 3 — अनुपालन मैट्रिक्स",
        "compliance_intro": "प्रत्येक निविदा आवश्यकता के अनुपालन की पुष्टि:",
        "h_requirement": "आवश्यकता", "h_response": "हमारा उत्तर", "h_status": "स्थिति",
        "sec_checklist": "खंड 4 — दस्तावेज़ सूची",
        "checklist_intro": "इस बोली के साथ संलग्न दस्तावेज़:",
        "h_sr": "क्र.", "h_document": "दस्तावेज़", "declaration": "घोषणा",
        "years": "वर्ष", "signatory_seal": "अधिकृत हस्ताक्षरकर्ता  |  मुहर",
    },
}


def build_docx(bid: dict, tender: dict, profile: dict, language: str = "en") -> bytes:
    """Compile bid JSON into a professional Word document. Returns raw bytes.

    language: 'en' or 'hi' — selects the fixed heading/label language. The AI body
    text inside `bid` is already in the language chosen at generation time.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    _lang = "hi" if str(language).lower().startswith("hi") else "en"
    _LBL  = _DOCX_LABELS[_lang]
    def lx(key: str) -> str:
        return _LBL.get(key, _DOCX_LABELS["en"].get(key, key))

    doc   = Document()
    today = date.today().strftime("%d %B %Y")
    cname = profile.get("company_name") or "Our Firm"

    NAVY  = RGBColor(0x1A, 0x52, 0x76)
    BLUE  = RGBColor(0x21, 0x62, 0x8E)

    def h(text: str, level: int = 1, para=None):
        p = para or doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16 - level * 2)
        r.font.color.rgb = NAVY if level == 1 else BLUE
        return p

    # ---- Cover Page ----
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(lx("bid_document"))
    r.bold, r.font.size, r.font.color.rgb = True, Pt(28), NAVY

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run(tender.get("title", "Tender Notice") or "Tender Notice").bold = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(
        f"\n{lx('tender_no')}: {tender.get('tender_no','—')}\n"
        f"{lx('issued_by')}: {tender.get('organization','—')}\n"
        f"{lx('submitted_by')}: {cname}\n"
        f"{lx('date')}: {today}"
    )
    doc.add_page_break()

    # ---- Cover Letter ----
    cl = bid.get("cover_letter", {})
    h(lx("cover_letter"), 1)
    doc.add_paragraph()
    doc.add_paragraph(f"{lx('date')}: {today}")
    doc.add_paragraph()
    doc.add_paragraph(cl.get("to") or lx("to_authority"))
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(cl.get("subject", "")).bold = True
    p = doc.add_paragraph()
    p.add_run(cl.get("ref", "")).italic = True
    doc.add_paragraph()
    doc.add_paragraph(cl.get("salutation") or lx("salutation"))
    doc.add_paragraph()
    for para in (cl.get("body_paragraphs") or []):
        doc.add_paragraph(para)
        doc.add_paragraph()
    if cl.get("closing"):
        doc.add_paragraph(cl["closing"])
    doc.add_paragraph()
    doc.add_paragraph(lx("yours_faithfully"))
    sig = doc.add_paragraph()
    sig.add_run(f"\n\n{'_' * 30}\n{cname}\n{lx('authorized_signatory')}\n{lx('date')}: {today}")
    doc.add_page_break()

    # ---- Company Profile ----
    h(lx("sec_company"), 1)
    cp_data = bid.get("company_profile", {})
    if cp_data.get("overview"):
        doc.add_paragraph(cp_data["overview"])
    doc.add_paragraph()

    h(lx("firm_details"), 2)
    # Defensive: these may arrive as text strings (DB / profile) — a format spec on
    # a string raises ValueError and would crash the whole bid build mid-way.
    try:
        _to_val = float(profile.get("turnover_lakhs") or 0)
        if _to_val != _to_val:  # NaN guard
            _to_val = 0.0
    except (TypeError, ValueError):
        _to_val = 0.0
    try:
        _exp_val = int(float(profile.get("experience_years") or 0))
    except (TypeError, ValueError):
        _exp_val = 0
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for label, val in [
        (lx("company_name"),     cname),
        (lx("contractor_class"), profile.get("contractor_class","—")),
        (lx("annual_turnover"),  f"Rs {_to_val:.0f} Lakhs"),
        (lx("experience"),       f"{_exp_val} {lx('years')}"),
        (lx("sectors"),          ", ".join(profile.get("sectors") or [])),
        (lx("states"),           ", ".join(profile.get("states") or [])),
    ]:
        row = tbl.add_row().cells
        row[0].text, row[1].text = label, str(val)
        if row[0].paragraphs[0].runs:
            row[0].paragraphs[0].runs[0].bold = True

    if cp_data.get("key_strengths"):
        doc.add_paragraph()
        h(lx("key_strengths"), 2)
        for s in cp_data["key_strengths"]:
            doc.add_paragraph(s, style="List Bullet")
    doc.add_page_break()

    # ---- Technical Proposal ----
    h(lx("sec_technical"), 1)
    tp = bid.get("technical_proposal", {})

    if tp.get("scope_understanding"):
        doc.add_paragraph()
        h(lx("t_scope"), 2)
        doc.add_paragraph(tp["scope_understanding"])

    if tp.get("methodology"):
        doc.add_paragraph()
        h(lx("t_method"), 2)
        for step in tp["methodology"]:
            doc.add_paragraph(step, style="List Bullet")

    if tp.get("team_structure"):
        doc.add_paragraph()
        h(lx("t_team"), 2)
        doc.add_paragraph(tp["team_structure"])

    if tp.get("quality_assurance"):
        doc.add_paragraph()
        h(lx("t_qa"), 2)
        doc.add_paragraph(tp["quality_assurance"])

    if tp.get("timeline_overview"):
        doc.add_paragraph()
        h(lx("t_timeline"), 2)
        doc.add_paragraph(tp["timeline_overview"])

    doc.add_page_break()

    # ---- Compliance Matrix ----
    h(lx("sec_compliance"), 1)
    doc.add_paragraph(lx("compliance_intro"))
    doc.add_paragraph()
    matrix = bid.get("compliance_matrix") or []
    if matrix:
        mtbl = doc.add_table(rows=1, cols=3)
        mtbl.style = "Table Grid"
        hdr = mtbl.rows[0].cells
        for i, txt in enumerate([lx("h_requirement"), lx("h_response"), lx("h_status")]):
            hdr[i].text = txt
            if hdr[i].paragraphs[0].runs:
                hdr[i].paragraphs[0].runs[0].bold = True
        for item in matrix:
            row = mtbl.add_row().cells
            row[0].text = item.get("requirement","")
            row[1].text = item.get("our_response","")
            row[2].text = item.get("status","")
    doc.add_page_break()

    # ---- Document Checklist ----
    h(lx("sec_checklist"), 1)
    doc.add_paragraph(lx("checklist_intro"))
    doc.add_paragraph()
    checklist = bid.get("document_checklist") or []
    if checklist:
        dtbl = doc.add_table(rows=1, cols=3)
        dtbl.style = "Table Grid"
        hdr  = dtbl.rows[0].cells
        for i, txt in enumerate([lx("h_sr"), lx("h_document"), lx("h_status")]):
            hdr[i].text = txt
            if hdr[i].paragraphs[0].runs:
                hdr[i].paragraphs[0].runs[0].bold = True
        for i, item in enumerate(checklist, 1):
            row = dtbl.add_row().cells
            row[0].text = str(i)
            row[1].text = item.get("document","")
            row[2].text = item.get("status","")
    doc.add_page_break()

    # ---- Declaration ----
    h(lx("declaration"), 1)
    doc.add_paragraph()
    doc.add_paragraph(bid.get("declaration","We declare that all information provided is true and correct."))
    doc.add_paragraph()
    doc.add_paragraph(f"{lx('date')}: {today}")
    doc.add_paragraph()
    doc.add_paragraph("_" * 30)
    doc.add_paragraph(cname)
    doc.add_paragraph(lx("signatory_seal"))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
